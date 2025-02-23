from typing import List, Tuple, Dict
import logging
from pathlib import Path
import io
from docx import Document as DocxDocument
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document
from .base import BaseDocumentLoader
from processors.image_extractor import ImageExtractor

logger = logging.getLogger(__name__)

class WordLoader(BaseDocumentLoader):
    """Word document loader - extracts text and embedded images"""
    
    def __init__(self, file_path: str):
        """Initialize loader"""
        super().__init__(file_path)
        self.image_extractor = ImageExtractor()
        self.image_map = {}  # Map to store image positions
    
    def _extract_images(self, docx_doc: DocxDocument) -> Dict[str, dict]:
        """
        Extract images from Word document and map them to their relationships
        Returns a dict mapping relationship IDs to image info
        """
        images = {}
        image_index = 1
        
        for rel in docx_doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    img_stream = io.BytesIO(image_data)
                    
                    # Process image
                    logger.info(f"Processing embedded image {image_index}")
                    extraction_result = self.image_extractor.extract_info(img_stream)
                    
                    if extraction_result["status"] == "success":
                        images[rel.rId] = {
                            "index": image_index,
                            "content": extraction_result["content"],
                            "status": "success"
                        }
                    else:
                        images[rel.rId] = {
                            "index": image_index,
                            "error": extraction_result.get("error", "Unknown error"),
                            "status": "failed"
                        }
                    
                    image_index += 1
                    
                except Exception as e:
                    logger.error(f"Failed to process image {image_index}: {str(e)}")
                    images[rel.rId] = {
                        "index": image_index,
                        "error": str(e),
                        "status": "failed"
                    }
                    image_index += 1
        
        return images
    
    def _process_paragraph(self, paragraph: Paragraph) -> str:
        """Process paragraph including any inline images"""
        text_parts = []
        
        for run in paragraph.runs:
            # Check for images in this run
            for element in run._element:
                if element.tag.endswith('drawing'):
                    # Find image relationship ID
                    rId = None
                    for child in element.iter():
                        if child.tag.endswith('blip'):
                            rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            break
                    
                    if rId and rId in self.image_map:
                        img = self.image_map[rId]
                        if img["status"] == "success":
                            text_parts.append(
                                f'\n<image id="{img["index"]:03d}">\n'
                                f'{img["content"]}\n'
                                f'</image>\n'
                            )
                        else:
                            text_parts.append(
                                f'\n<image id="{img["index"]:03d}" status="failed">\n'
                                f'Image processing failed: {img.get("error", "Unknown error")}\n'
                                f'</image>\n'
                            )
            
            text_parts.append(run.text)
        
        return "".join(text_parts)
    
    def _process_table(self, table: Table) -> str:
        """Process table content"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = []
                for paragraph in cell.paragraphs:
                    text = self._process_paragraph(paragraph)
                    if text.strip():
                        cell_text.append(text)
                cells.append(" ".join(cell_text))
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def _iter_block_items(self, parent):
        """Iterate through all blocks (paragraphs and tables)"""
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Something's not right")
        
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
    
    def _extract_content(self, docx_doc: DocxDocument) -> Tuple[str, List[dict]]:
        """Extract text content and process images"""
        content_parts = []
        self.image_map = self._extract_images(docx_doc)
        
        # Process all blocks (paragraphs and tables)
        for block in self._iter_block_items(docx_doc):
            if isinstance(block, Paragraph):
                text = self._process_paragraph(block)
                if text.strip():
                    content_parts.append(text)
            elif isinstance(block, Table):
                table_text = self._process_table(block)
                if table_text.strip():
                    content_parts.append(table_text)
        
        # Convert image map to list for metadata
        images = [img for img in self.image_map.values()]
        return "\n\n".join(content_parts), images
    
    def load(self) -> List[Document]:
        """Load Word document and extract content"""
        try:
            # Check file exists
            doc_path = Path(self.file_path)
            if not doc_path.exists():
                raise FileNotFoundError(f"Word document not found: {doc_path}")
            
            # Load document
            logger.info(f"Loading Word document: {doc_path}")
            docx_doc = DocxDocument(doc_path)
            
            # Extract content
            content, images = self._extract_content(docx_doc)
            
            # Create Document object
            doc = Document(
                page_content=content,
                metadata={
                    "source": str(doc_path),
                    "file_type": "word",
                    "file_name": doc_path.name,
                    "images": images
                }
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"Error loading Word document: {str(e)}", exc_info=True)
            # Return error document
            error_doc = Document(
                page_content=f"Failed to load Word document: {str(e)}",
                metadata={
                    "source": str(self.file_path),
                    "file_type": "word",
                    "extraction_status": "failed",
                    "error": str(e)
                }
            )
            return [error_doc] 